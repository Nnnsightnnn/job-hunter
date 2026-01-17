"""
Resume Parser Module
Handles text extraction from uploaded resume files (PDF, DOCX, TXT, JSON)
"""
import json
import os
import mimetypes
from pathlib import Path
from typing import Tuple, Optional

# PDF parsing
try:
    import pdfplumber
    PDFPLUMBER_AVAILABLE = True
except ImportError:
    PDFPLUMBER_AVAILABLE = False

try:
    import PyPDF2
    PYPDF2_AVAILABLE = True
except ImportError:
    PYPDF2_AVAILABLE = False

# DOCX parsing
try:
    from docx import Document
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False


ALLOWED_EXTENSIONS = {'pdf', 'docx', 'txt', 'json'}
MAX_FILE_SIZE = 5 * 1024 * 1024  # 5MB


class ResumeParser:
    """
    Handles file validation and text extraction from various resume formats
    """

    def __init__(self, upload_folder: str = "data/uploads"):
        self.upload_folder = Path(upload_folder)
        self.upload_folder.mkdir(parents=True, exist_ok=True)

    def validate_file(self, filename: str, file_size: int) -> Tuple[bool, str]:
        """
        Validate uploaded file by extension, size, and MIME type

        Returns:
            Tuple of (is_valid, error_message)
        """
        if not filename:
            return False, "No file provided"

        # Check extension
        ext = filename.rsplit('.', 1)[-1].lower() if '.' in filename else ''
        if ext not in ALLOWED_EXTENSIONS:
            return False, f"Please upload a PDF, DOCX, TXT, or JSON file"

        # Check size
        if file_size > MAX_FILE_SIZE:
            return False, f"File exceeds 5MB limit"

        # Check MIME type
        mime_type, _ = mimetypes.guess_type(filename)
        valid_mimes = {
            'application/pdf',
            'application/vnd.openxmlformats-officedocument.wordprocessingml.document',
            'text/plain',
            'application/json'
        }

        # MIME check is advisory - extension check is primary
        if mime_type and mime_type not in valid_mimes:
            # Some systems may not properly detect MIME, so we allow if extension is valid
            pass

        return True, ""

    def get_file_extension(self, filename: str) -> str:
        """Get the lowercase file extension"""
        return filename.rsplit('.', 1)[-1].lower() if '.' in filename else ''

    def extract_text(self, filepath: str) -> Tuple[str, Optional[str]]:
        """
        Extract text from a file based on its extension

        Returns:
            Tuple of (extracted_text, error_message)
        """
        path = Path(filepath)
        ext = path.suffix.lower().lstrip('.')

        extractors = {
            'pdf': self.extract_text_from_pdf,
            'docx': self.extract_text_from_docx,
            'txt': self.extract_text_from_txt,
            'json': self.extract_text_from_json
        }

        extractor = extractors.get(ext)
        if not extractor:
            return "", f"Unsupported file format: {ext}"

        return extractor(filepath)

    def extract_text_from_pdf(self, filepath: str) -> Tuple[str, Optional[str]]:
        """
        Extract text from PDF using pdfplumber (primary) or PyPDF2 (fallback)
        """
        text = ""
        error = None

        # Try pdfplumber first (better layout handling)
        if PDFPLUMBER_AVAILABLE:
            try:
                with pdfplumber.open(filepath) as pdf:
                    pages = []
                    for page in pdf.pages:
                        page_text = page.extract_text()
                        if page_text:
                            pages.append(page_text)
                    text = "\n\n".join(pages)

                if text.strip():
                    return text.strip(), None
            except Exception as e:
                error = f"pdfplumber error: {str(e)}"

        # Fallback to PyPDF2
        if PYPDF2_AVAILABLE:
            try:
                with open(filepath, 'rb') as f:
                    reader = PyPDF2.PdfReader(f)
                    pages = []
                    for page in reader.pages:
                        page_text = page.extract_text()
                        if page_text:
                            pages.append(page_text)
                    text = "\n\n".join(pages)

                if text.strip():
                    return text.strip(), None
            except Exception as e:
                error = f"PyPDF2 error: {str(e)}"

        if not text.strip():
            return "", "Could not extract text (may be image-based or protected)"

        return text.strip(), error

    def extract_text_from_docx(self, filepath: str) -> Tuple[str, Optional[str]]:
        """
        Extract text from DOCX using python-docx
        """
        if not DOCX_AVAILABLE:
            return "", "DOCX support not available. Install python-docx."

        try:
            doc = Document(filepath)
            paragraphs = []

            for para in doc.paragraphs:
                if para.text.strip():
                    paragraphs.append(para.text.strip())

            # Also extract from tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        if cell.text.strip():
                            row_text.append(cell.text.strip())
                    if row_text:
                        paragraphs.append(" | ".join(row_text))

            text = "\n".join(paragraphs)

            if not text.strip():
                return "", "Document appears to be empty"

            return text.strip(), None

        except Exception as e:
            return "", f"Error reading DOCX: {str(e)}"

    def extract_text_from_txt(self, filepath: str) -> Tuple[str, Optional[str]]:
        """
        Extract text from TXT with encoding detection
        """
        encodings = ['utf-8', 'utf-16', 'latin-1', 'cp1252']

        for encoding in encodings:
            try:
                with open(filepath, 'r', encoding=encoding) as f:
                    text = f.read()
                    if text.strip():
                        return text.strip(), None
            except (UnicodeDecodeError, UnicodeError):
                continue
            except Exception as e:
                return "", f"Error reading file: {str(e)}"

        return "", "Could not determine file encoding"

    def extract_text_from_json(self, filepath: str) -> Tuple[str, Optional[str]]:
        """
        Parse JSON resume and convert to text for display, or validate structure
        """
        try:
            with open(filepath, 'r', encoding='utf-8') as f:
                data = json.load(f)

            # Return the raw JSON as pretty-printed string for preview
            # The actual structured data will be used directly
            return json.dumps(data, indent=2), None

        except json.JSONDecodeError as e:
            return "", f"Invalid JSON format: {str(e)}"
        except Exception as e:
            return "", f"Error reading JSON: {str(e)}"

    def parse_json_resume(self, filepath: str) -> Tuple[Optional[dict], Optional[str]]:
        """
        Load and validate JSON resume format

        Returns:
            Tuple of (parsed_data, error_message)
        """
        try:
            with open(filepath, 'r', encoding='utf-8') as f:
                data = json.load(f)

            # Basic validation - check for expected sections
            expected_keys = ['personal', 'experience', 'education', 'skills']
            missing = [k for k in expected_keys if k not in data]

            if missing:
                # Not in our expected format, but still valid JSON
                # Let the structurer handle conversion
                return data, f"JSON loaded, but missing sections: {', '.join(missing)}"

            return data, None

        except json.JSONDecodeError as e:
            return None, f"Invalid JSON format: {str(e)}"
        except Exception as e:
            return None, f"Error parsing JSON: {str(e)}"

    def cleanup_temp_file(self, filepath: str) -> bool:
        """
        Remove uploaded file after processing
        """
        try:
            path = Path(filepath)
            if path.exists():
                path.unlink()
            return True
        except Exception:
            return False
